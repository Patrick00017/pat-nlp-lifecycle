import logging
import shutil
from pathlib import Path

import fitz
from docx import Document as DocxDocument
from langchain_community.vectorstores import Chroma
from langchain_core.documents import Document
from langchain_huggingface import HuggingFaceEmbeddings
from langchain_text_splitters import RecursiveCharacterTextSplitter

logger = logging.getLogger("rag-service")

SANDBOX_ROOT = Path(__file__).parent
DEFAULT_PDF = SANDBOX_ROOT / "assets" / "usermanual.pdf"
MODEL_PATH = SANDBOX_ROOT / "model" / "bge-small-zh-v1.5"
CHROMA_DIR = SANDBOX_ROOT / "chroma_db"
UPLOAD_DIR = SANDBOX_ROOT / "uploads"

LOADERS = {}


def _resolve_default_files(file_paths):
    if file_paths:
        return file_paths
    if DEFAULT_PDF.exists():
        return [str(DEFAULT_PDF)]
    return []


def _load_txt(file_path: str) -> list[Document]:
    name = Path(file_path).name
    with open(file_path, "r", encoding="utf-8") as f:
        text = f.read()
    if text.strip():
        return [Document(page_content=text, metadata={"source": name})]
    return []


def _load_pdf_fitz(file_path: str) -> list[Document]:
    name = Path(file_path).name
    docs = []
    with fitz.open(file_path) as pdf:
        for page_num in range(len(pdf)):
            page = pdf[page_num]
            text = page.get_text()
            if text.strip():
                docs.append(Document(
                    page_content=text,
                    metadata={"source": name, "page": page_num + 1},
                ))
    return docs


def _load_docx(file_path: str) -> list[Document]:
    name = Path(file_path).name
    doc = DocxDocument(file_path)
    paragraphs = [p.text for p in doc.paragraphs if p.text.strip()]
    text = "\n".join(paragraphs)
    if text.strip():
        return [Document(page_content=text, metadata={"source": name})]
    return []


EXT_LOADER = {
    ".txt": _load_txt,
    ".md": _load_txt,
    ".pdf": _load_pdf_fitz,
    ".docx": _load_docx,
}


def load_documents(file_paths: list[str]) -> list[Document]:
    all_docs = []
    for fp in file_paths:
        ext = Path(fp).suffix.lower()
        loader = EXT_LOADER.get(ext)
        if not loader:
            logger.warning("unsupported file type: %s", fp)
            continue
        try:
            docs = loader(fp)
            logger.info("loaded %d docs from %s", len(docs), fp)
            all_docs.extend(docs)
        except Exception:
            logger.exception("failed to load %s", fp)
    return all_docs


def build_chroma(file_paths=None, force_rebuild=False):
    embeddings = HuggingFaceEmbeddings(
        model_name=str(MODEL_PATH),
        model_kwargs={"device": "cpu"},
        encode_kwargs={"normalize_embeddings": True},
    )

    if force_rebuild and CHROMA_DIR.exists():
        logger.info("force rebuild: removing %s", CHROMA_DIR)
        shutil.rmtree(str(CHROMA_DIR))

    resolved = _resolve_default_files(file_paths)

    if CHROMA_DIR.exists():
        logger.info("loading existing Chroma from %s", CHROMA_DIR)
        vs = Chroma(persist_directory=str(CHROMA_DIR), embedding_function=embeddings)
        return vs, vs._collection.count()

    if not resolved:
        logger.warning("no source files found, returning empty Chroma")
        vs = Chroma.from_documents(
            documents=[Document(page_content="placeholder")],
            embedding=embeddings,
            persist_directory=str(CHROMA_DIR),
        )
        return vs, vs._collection.count()

    logger.info("loading %d source files: %s", len(resolved), resolved)
    docs = load_documents(resolved)
    logger.info("loaded %d documents total, chunking ...", len(docs))
    splitter = RecursiveCharacterTextSplitter(chunk_size=500, chunk_overlap=100)
    chunks = splitter.split_documents(docs)
    logger.info("created %d chunks, indexing into Chroma ...", len(chunks))
    vs = Chroma.from_documents(
        documents=chunks,
        embedding=embeddings,
        persist_directory=str(CHROMA_DIR),
    )
    return vs, vs._collection.count()
